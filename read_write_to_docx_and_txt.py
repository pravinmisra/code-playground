# -*- coding: utf-8 -*-
"""
Created on Sun May 25 13:56:36 2025

@author: Admin
"""

from docx import Document

#read an existing text document

with open("C:\\code-playground\\test_data\\test_data_1.txt") as f:
    print("Initial text in the document ==> ")
    print(f.read())
    
#Append text to an existing text document at its end

with open("C:\\code-playground\\test_data\\test_data_1.txt","a") as f:
    f.write("A line appended at the end of the document")
    
#read the text document after appending new text
with open("C:\\code-playground\\test_data\\test_data_1.txt") as f:
    print("After appending ==> ")
    print(f.read())
    
#replace text to an existing text document at its end

with open("C:\\code-playground\\test_data\\test_data_1.txt","w") as f:
    f.write("The text document has been overwritten by this new line")
    
#read the text document after replacement
with open("C:\\code-playground\\test_data\\test_data_1.txt") as f:
    print("After overwriting ==> ")
    print(f.read())


#read an existing word document

document = Document("C:\\code-playground\\test_data\\test_data_2.docx")

for paragraph in document.paragraphs:
    print(paragraph.text)
    
#create a new document and write in it

document_ = Document()

document_.add_paragraph("This a new document created")
document_.add_paragraph("This is second paragraph")
document_.add_paragraph("This is third paragraph")

document_.save("C:\\code-playground\\test_data\\test_data_3.docx")

#read the newly written document

for paragraph in document_.paragraphs:
    print(paragraph.text)



    
    



  
